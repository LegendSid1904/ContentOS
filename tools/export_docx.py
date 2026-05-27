"""
Export tool — converts text content to DOCX file.

Usage:
    python export_docx.py --content "Your content here..." --output output.docx
    python export_docx.py --input input.txt --output output.docx --title "My Document"

Dependencies:
    pip install python-docx
"""

import os
import argparse

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(
    content: str,
    output_path: str,
    title: str = "ContentOS Export",
) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para_text in content.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        if para_text.startswith("## "):
            doc.add_heading(para_text[3:], level=2)
        elif para_text.startswith("# "):
            doc.add_heading(para_text[2:], level=1)
        else:
            doc.add_paragraph(para_text)

    doc.save(output_path)
    return os.path.abspath(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Export content to DOCX")
    parser.add_argument("--content", help="Text content to export")
    parser.add_argument("--input", help="Input file path")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    parser.add_argument("--title", default="ContentOS Export")

    args = parser.parse_args()

    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            content = f.read()
    elif args.content:
        content = args.content
    else:
        raise ValueError("Provide either --content or --input")

    path = export_docx(content=content, output_path=args.output, title=args.title)
    print(f"DOCX exported to: {path}")
